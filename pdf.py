from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber

old_pdf = "input.pdf"
new_doc = "new.docx"

# Step 1: Use pdf2docx for the first page
obj = Converter(old_pdf)
obj.convert(new_doc, start=0, end=1)
obj.close()

# Step 2: Use pdfplumber for the remaining pages
doc = Document(new_doc)

with pdfplumber.open(old_pdf) as pdf:
    for i, page in enumerate(pdf.pages[1:], start=1):  # Start from the second page
        text = page.extract_text()
        if text:
            for line in text.split('\n'):
                paragraph = doc.add_paragraph(line)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.runs
                for r in run:
                    r.font.size = Pt(12)
        
        tables = page.extract_tables()
        for table in tables:
            table_doc = doc.add_table(rows=len(table), cols=len(table[0]))
            for row_idx, row in enumerate(table):
                for col_idx, cell in enumerate(row):
                    table_doc.cell(row_idx, col_idx).text = cell
                    table_doc.cell(row_idx, col_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(new_doc)




